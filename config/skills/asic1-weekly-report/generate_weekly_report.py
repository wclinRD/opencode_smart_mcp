#!/usr/bin/env python3
"""
ASIC1 Weekly Report Generator

Template-based. Copies existing DennisLin docx, replaces content in-place.
Preserves all original formatting: numbered list style (a9), theme fonts,
multi-level indentation, paragraph spacing, alignment, and text colors.
"""

import argparse
import json
import sys
import os
import shutil
import copy

try:
    import docx
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("Error: Missing python-docx. Run: pip3 install python-docx")
    sys.exit(1)


W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _q(tag):
    return f'{{{W}}}{tag}'


SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.join(SKILL_DIR, "template.docx")


def find_template():
    candidates = [
        DEFAULT_TEMPLATE,
        os.path.expanduser("~/Downloads"),
        os.path.expanduser(
            "~/Library/CloudStorage/CloudMounter-pCloud/Project/Meeting/2026"
        ),
    ]
    for base in candidates:
        if not os.path.isdir(base) and not os.path.isfile(base):
            continue
        if os.path.isfile(base):
            return base
        for f in sorted(os.listdir(base), reverse=True):
            if f.startswith("DennisLin_1_2_ZA") and f.endswith(".docx"):
                if not f.startswith("~$"):
                    return os.path.join(base, f)
    return None


# ──────────────────────────────────────────────
# Header editing (Table 0)
# ──────────────────────────────────────────────

def set_cell_text_by_run(cell, text):
    """
    Replace text in the first run of cell's first paragraph.
    Preserves all formatting. Removes extra runs.
    """
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    runs = list(p.runs)
    if runs:
        runs[0].text = str(text)
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        run = p.add_run(str(text))


# ──────────────────────────────────────────────
# Nested table editing (Table 1)
# ──────────────────────────────────────────────

def _find_section_header(cell):
    """
    Find a section header paragraph: no pStyle, no numPr, has bold.
    Returns None if not found (fallback to numbered template).
    """
    for p in cell.findall(_q('p')):
        pPr = p.find(_q('pPr'))
        if pPr is None:
            continue
        # Has pStyle or numPr? Skip (these are numbered items)
        if pPr.find(_q('pStyle')) is not None:
            continue
        if pPr.find(_q('numPr')) is not None:
            continue
        # Has bold?
        rPr = p.find(f'.//{_q("rPr")}')
        if rPr is not None and rPr.find(_q('b')) is not None:
            return p
    # Fallback: any paragraph without pStyle
    for p in cell.findall(_q('p')):
        pPr = p.find(_q('pPr'))
        if pPr is None:
            continue
        if pPr.find(_q('pStyle')) is not None:
            continue
        if pPr.find(_q('numPr')) is not None:
            continue
        if p.findall(_q('r')):
            return p
    return None


def _remove_numpr(para_elem):
    """Remove numPr from paragraph if present."""
    pPr = para_elem.find(_q('pPr'))
    if pPr is not None:
        numPr = pPr.find(_q('numPr'))
        if numPr is not None:
            pPr.remove(numPr)


def _remove_indent(para_elem):
    """Remove indentation from paragraph if present."""
    pPr = para_elem.find(_q('pPr'))
    if pPr is not None:
        ind = pPr.find(_q('ind'))
        if ind is not None:
            pPr.remove(ind)


def _ensure_bold(para_elem):
    """Ensure first run has bold formatting."""
    for rPr in para_elem.findall(f'.//{_q("rPr")}'):
        b = rPr.find(_q('b'))
        if b is None:
            b = etree.SubElement(rPr, _q('b'))
            rPr.append(b)
        # Set val if it doesn't exist (has effect without val too)
        if _q('val') not in b.attrib:
            b.set(_q('val'), '1')


def update_nested_table(nested_tbl, items):
    """
    Update nested table content while preserving all formatting.
    
    Strategy:
    1. Find the data row (Row 1, two cells: items | progress)
    2. Find a "good template" paragraph (has numId=11 for items,
       has jc=center for progress)
    3. Remove ALL existing content paragraphs
    4. Clone good template N times and set text
    5. Apply gray color (A6A6A6) to items with "Done" progress
    """
    rows = nested_tbl.findall(_q('tr'))
    if len(rows) < 2:
        print("⚠️  Nested table has < 2 rows, skipping.")
        return
    
    data_row = rows[1]
    cells = data_row.findall(_q('tc'))
    if len(cells) < 2:
        print("⚠️  Data row has < 2 cells.")
        return
    
    items_cell = cells[0]
    progress_cell = cells[1]
    
    # ── Find template paragraphs ──
    # Numbered template: has numId=11 (for items with level >= 0)
    item_tpl = _find_template(items_cell, _q('numId'), '11')
    # Section header template: NO pStyle, NO numPr, has bold (for level=-1)
    header_tpl = _find_section_header(items_cell)
    
    if item_tpl is None:
        print("⚠️  No numbered template found in items cell.")
        return
    
    # Progress template: center-aligned with text content
    prog_tpl = _find_template(progress_cell, _q('jc'), 'center', require_text=True)
    if prog_tpl is None:
        for p in progress_cell.findall(_q('p')):
            if p.findall(_q('r')):
                prog_tpl = p
                break
    if prog_tpl is None:
        print("⚠️  No progress template found.")
        return
    
    n_needed = len(items)
    
    # ── Remove ALL existing paragraphs ──
    _clear_cell(items_cell)
    _clear_cell(progress_cell)
    
    # ── Clone templates N times ──
    for i, item_entry in enumerate(items):
        # Support both tuple and dict with optional 'level'
        if isinstance(item_entry, dict):
            item_text = item_entry.get("item", "")
            progress_text = item_entry.get("progress", "")
            level = item_entry.get("level", 0)
        elif isinstance(item_entry, (tuple, list)):
            item_text = item_entry[0]
            progress_text = item_entry[1] if len(item_entry) > 1 else ""
            level = item_entry[2] if len(item_entry) > 2 else 0
        else:
            item_text = str(item_entry)
            progress_text = ""
            level = 0
        
        # Choose template: header template for level=-1, numbered for level>=0
        is_header = (int(level) < 0)
        if is_header and header_tpl is not None:
            new_item_p = copy.deepcopy(header_tpl)
        else:
            new_item_p = copy.deepcopy(item_tpl)
        new_prog_p = copy.deepcopy(prog_tpl)
        
        # Set numbering (header template has no numPr, numbered does)
        if not is_header:
            _set_numbering_level(new_item_p, int(level))
        else:
            # Section header: ensure no numPr and no indentation
            _remove_numpr(new_item_p)
            _remove_indent(new_item_p)
            # Ensure bold is present (header template should have it)
            _ensure_bold(new_item_p)
        
        # Ensure progress paragraph has negative indent (center alignment)
        _ensure_progress_indent(new_prog_p)
        
        # Set text
        _set_para_text(new_item_p, item_text)
        _set_para_text(new_prog_p, progress_text)
        
        # Apply gray color for Done items
        progress_lower = str(progress_text).strip().lower()
        if progress_lower == 'done':
            _apply_gray_color(new_item_p)
        else:
            _remove_gray_color(new_item_p)
        
        items_cell.append(new_item_p)
        progress_cell.append(new_prog_p)
    
    # ── Set row height ──
    _set_row_height(data_row, 4210)


def _find_template(cell, attr_tag, attr_val, require_text=False):
    """
    Find first paragraph in cell that has an element with given value.
    e.g. _find_template(cell, 'numId', '11') finds first para with numId=11.
    If require_text=True, only matches paragraphs with at least one <w:r>.
    """
    for p in cell.findall(_q('p')):
        # Check for text content if required
        if require_text:
            runs = p.findall(_q('r'))
            if not runs:
                continue
        # Search for element with matching attribute value
        for el in p.iter():
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            target = attr_tag.split('}')[-1] if '}' in attr_tag else attr_tag
            if tag == target:
                if el.get(_q('val')) == attr_val:
                    return p
    return None


def _clear_cell(cell):
    """Remove all paragraphs from a cell."""
    for p in cell.findall(_q('p')):
        cell.remove(p)


def _set_para_text(para_elem, text):
    """
    Set paragraph text by replacing content of first <w:t> element.
    Preserves all formatting, numbering, styles.
    Removes all but first <w:r> to avoid extra runs.
    Removes extra <w:t> elements inside first run.
    If text is empty, preserves <w:t> element but sets empty text.
    If no <w:t> found and text is required, creates one.
    """
    runs = para_elem.findall(_q('r'))
    
    if not runs:
        # No runs at all - create one with a t element
        if str(text).strip():
            new_run = etree.SubElement(para_elem, _q('r'))
            # Copy rPr from any existing source if available
            new_t = etree.SubElement(new_run, _q('t'))
            new_t.text = str(text)
        return
    
    # Remove extra runs beyond the first
    for r in runs[1:]:
        para_elem.remove(r)
    
    first_run = runs[0]
    t_list = first_run.findall(_q('t'))
    
    if t_list:
        t_elem = t_list[0]
        t_elem.text = str(text) if str(text) else None
        # Remove extra t elements
        for t in t_list[1:]:
            first_run.remove(t)
    else:
        # No t element in first run - create one
        if str(text).strip():
            new_t = etree.SubElement(first_run, _q('t'))
            new_t.text = str(text)


def _apply_gray_color(para_elem):
    """Apply gray color (A6A6A6) to all runs in a paragraph."""
    for rPr in para_elem.findall(f'.//{_q("rPr")}'):
        color = rPr.find(_q('color'))
        if color is not None:
            color.set(_q('val'), 'A6A6A6')
            color.set(_q('themeColor'), 'background1')
            color.set(_q('themeShade'), 'A6')


def _remove_gray_color(para_elem):
    """Remove gray color, restore default (000000)."""
    for rPr in para_elem.findall(f'.//{_q("rPr")}'):
        color = rPr.find(_q('color'))
        if color is not None and color.get(_q('val')) == 'A6A6A6':
            color.set(_q('val'), '000000')
            # Remove theme attributes if they were for gray
            for attr in [_q('themeColor'), _q('themeShade')]:
                if attr in color.attrib:
                    del color.attrib[attr]


def _set_numbering_level(para_elem, level):
    """
    Set the ilvl on a paragraph's numbering properties.
    If level < 0, remove entire numPr from pPr (for section headers).
    If level >= 0, ensure numPr exists and set ilvl to level.
    """
    pPr = para_elem.find(_q('pPr'))
    if pPr is None:
        return
    numPr = pPr.find(_q('numPr'))
    
    if level < 0:
        # Section header: remove numbering entirely
        if numPr is not None:
            pPr.remove(numPr)
        return
    
    # Ensure numPr exists (should, if cloned from numbered template)
    if numPr is None:
        numPr = etree.SubElement(pPr, _q('numPr'))
        pPr.append(numPr)
    
    # Set or create ilvl element
    ilvl = numPr.find(_q('ilvl'))
    if ilvl is None:
        ilvl = etree.SubElement(numPr, _q('ilvl'))
        numPr.append(ilvl)
    ilvl.set(_q('val'), str(level))
    
    # Ensure numId exists
    numId = numPr.find(_q('numId'))
    if numId is None:
        numId = etree.SubElement(numPr, _q('numId'))
        numPr.append(numId)
        numId.set(_q('val'), '11')


def _ensure_progress_indent(para_elem):
    """
    Ensure progress paragraph has negative indent for center alignment.
    Original uses: w:ind w:left="-79" w:leftChars="-33"
    """
    pPr = para_elem.find(_q('pPr'))
    if pPr is None:
        pPr = etree.SubElement(para_elem, _q('pPr'))
        para_elem.insert(0, pPr)
    ind = pPr.find(_q('ind'))
    if ind is None:
        ind = etree.SubElement(pPr, _q('ind'))
        pPr.append(ind)
    ind.set(_q('left'), '-79')
    ind.set(_q('leftChars'), '-33')


def _set_row_height(row, height):
    """Set row height (trHeight) for a table row."""
    trPr = row.find(_q('trPr'))
    if trPr is None:
        trPr = etree.SubElement(row, _q('trPr'))
        row.insert(0, trPr)
    trHeight = trPr.find(_q('trHeight'))
    if trHeight is None:
        trHeight = etree.SubElement(trPr, _q('trHeight'))
    trHeight.set(_q('val'), str(height))
    trHeight.set(_q('hRule'), 'atLeast')


# ──────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────

def generate_report(output_path, date, author, department="IC設計部",
                    items=None, template_path=None):
    if template_path is None:
        template_path = find_template()
    if not template_path or not os.path.exists(template_path):
        print(f"❌ No template found.")
        sys.exit(1)
    
    print(f"📋 Template: {template_path}")
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    
    # ── Table 0: Header ──
    t0 = doc.tables[0]
    set_cell_text_by_run(t0.rows[0].cells[1], department)
    set_cell_text_by_run(t0.rows[0].cells[4], date)
    set_cell_text_by_run(t0.rows[1].cells[4], author)
    
    # ── Table 1: Nested task table ──
    wrapper_cell = doc.tables[1].rows[0].cells[0]
    nested_tbl = wrapper_cell._tc.find(_q('tbl'))
    
    if nested_tbl is not None and items:
        update_nested_table(nested_tbl, items)
    
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

def parse_items_from_args(items_args):
    result = []
    for arg in items_args:
        if "::" in arg:
            parts = arg.split("::", 2)
            item = parts[0].strip()
            progress = parts[1].strip() if len(parts) > 1 else ""
            level = int(parts[2].strip()) if len(parts) > 2 else 0
            result.append({"item": item, "progress": progress, "level": level})
        else:
            result.append({"item": arg.strip(), "progress": "", "level": 0})
    return result


def parse_items_from_json(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    dept = data.get("department", "IC設計部")
    date = data.get("date", "")
    author = data.get("author", "")
    items = data.get("items", [])
    # Support both simple format (item+progress) and extended (item+progress+level)
    parsed_items = []
    for it in items:
        if isinstance(it, dict):
            parsed_items.append(it)
        else:
            parsed_items.append({"item": it[0], "progress": it[1] if len(it) > 1 else "", "level": it[2] if len(it) > 2 else 0})
    return dept, date, author, parsed_items


def main():
    parser = argparse.ArgumentParser(
        description="ASIC1 Weekly Report Generator"
    )
    parser.add_argument("--date", default="", help="日期 e.g. 2026/05/28")
    parser.add_argument("--author", default="林文仲", help="撰寫人")
    parser.add_argument("--department", default="IC設計部", help="部門")
    parser.add_argument("--items", nargs="*", help="'項目::進度'")
    parser.add_argument("--json-input", help="JSON 輸入檔")
    parser.add_argument("--template", default=None,
                        help="範本 DOCX 路徑 (預設: skill 目錄下 template.docx)")
    parser.add_argument("--output", default="weekly_report.docx",
                        help="輸出檔名")
    
    args = parser.parse_args()
    
    if args.json_input:
        dept, date, author, items = parse_items_from_json(args.json_input)
    else:
        dept = args.department
        date = args.date
        author = args.author
        items = parse_items_from_args(args.items or [])
    
    if not date:
        print("❌ 需要 --date 或 --json-input")
        sys.exit(1)
    
    template_path = args.template or DEFAULT_TEMPLATE
    if not os.path.exists(template_path):
        alt = find_template()
        if alt:
            print(f"  改用: {alt}")
            template_path = alt
        else:
            print("❌ 找不到範本")
            sys.exit(1)
    args.template = template_path
    
    generate_report(
        output_path=args.output,
        date=date,
        author=author,
        department=dept,
        items=items,
        template_path=args.template,
    )


if __name__ == "__main__":
    main()
