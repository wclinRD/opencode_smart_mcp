from __future__ import annotations

import zipfile
import sys
from ..exceptions import ExtractionError


def extract_docx_with_python_docx(docx_path: str) -> str | None:
    try:
        import docx
        document = docx.Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        return None
    except Exception as e:
        print(f"  [warn] extract_docx_with_python_docx failed: {type(e).__name__}: {e}", file=sys.stderr)
        return None


def extract_docx_with_zipfile(docx_path: str) -> str | None:
    try:
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(docx_path) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts: list[str] = []
        for paragraph in root.iter(f"{namespace}p"):
            texts = [node.text for node in paragraph.iter(f"{namespace}t") if node.text]
            if texts:
                parts.append("".join(texts))
        return "\n".join(parts) if parts else None
    except Exception as e:
        print(f"  [warn] extract_docx_with_zipfile failed: {type(e).__name__}: {e}", file=sys.stderr)
        return None


def validate_docx_xml_safety(docx_path: str) -> None:
    """Scan all XML files in the DOCX zip archive to prevent XML Entity Expansion (Billion Laughs) and XXE injections."""
    try:
        with zipfile.ZipFile(docx_path) as zf:
            for name in zf.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    xml_bytes = zf.read(name)
                    for encoding in ("utf-8", "utf-16", "utf-16le", "utf-16be", "utf-32"):
                        try:
                            content = xml_bytes.decode(encoding, errors="ignore").upper()
                        except LookupError:
                            continue
                        if "<!DOCTYPE" in content or "<!ENTITY" in content:
                            raise ExtractionError(
                                f"Security validation failed: XML file '{name}' in DOCX archive contains forbidden DTD or entity declarations."
                            )
    except zipfile.BadZipFile as e:
        raise ExtractionError(f"Invalid DOCX file: {e}")
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Error during security validation of DOCX archive: {e}")


def extract_docx(docx_path: str) -> tuple[str, str]:
    validate_docx_xml_safety(docx_path)
    print("Trying python-docx...", end=" ", flush=True)
    text = extract_docx_with_python_docx(docx_path)
    if text and text.strip():
        print("OK")
        return text, "python-docx"

    print("not available")
    print("Trying stdlib DOCX parser...", end=" ", flush=True)
    text = extract_docx_with_zipfile(docx_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile-docx"

    print("FAILED")
    raise ExtractionError(
        "Could not extract text from DOCX.\n"
        "Install python-docx for best results:\n"
        "  pip3 install python-docx"
    )
